import flet as ft
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_center_text(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_left_text(doc, text, size=13, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def export_docx(e):
    doc = Document()

    # ===== QUỐC HIỆU =====
    add_center_text(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", 14, True)
    add_center_text(doc, "Độc lập – Tự do – Hạnh phúc", 12)
    add_center_text(doc, "———o0o———", 12)

    # ===== TIÊU ĐỀ =====
    add_center_text(doc, "HỢP ĐỒNG THUÊ XE", 16, True)
    add_center_text(doc, "Số: ……/2026/HĐTX-PiCar", 12)

    add_left_text(doc, "- Căn cứ Bộ Luật Dân sự 2015;")
    add_left_text(doc, "- Căn cứ Luật Thương mại 2005;")
    add_left_text(doc, "- Căn cứ vào nhu cầu và khả năng của các bên.")

    add_left_text(
        doc,
        "\nHôm nay, ngày .... tháng .... năm 20..., tại văn phòng PiCar, chúng tôi gồm:",
    )

    # ===== BÊN A =====
    add_left_text(doc, "\nBÊN A: BÊN CHO THUÊ (CÔNG TY PiCar)", bold=True)
    add_left_text(doc, "- Địa chỉ: Tầng 1, Tòa nhà Tech, TP.HCM")
    add_left_text(doc, "- Đại diện: Ông Nguyễn Văn A")
    add_left_text(doc, "- Chức vụ: Giám đốc điều hành")

    # ===== BÊN B =====
    add_left_text(doc, "\nBÊN B: BÊN THUÊ (KHÁCH HÀNG)", bold=True)
    add_left_text(doc, "- Ông/Bà: ................................................")
    add_left_text(doc, "- Số CCCD/CMND: .........................................")
    add_left_text(doc, "- Địa chỉ thường trú: .....................................")
    add_left_text(doc, "- Số điện thoại: ........................................")

    doc.save("Hop_dong_thue_xe_PiCar.docx")

    e.page.snack_bar = ft.SnackBar(
        ft.Text("✅ Đã xuất file Hop_dong_thue_xe_PiCar.docx")
    )
    e.page.snack_bar.open = True
    e.page.update()


def main(page: ft.Page):
    page.title = "Hợp đồng thuê xe - PiCar"
    page.bgcolor = ft.Colors.WHITE
    page.scroll = ft.ScrollMode.AUTO

    def title(text, size=16):
        return ft.Text(
            text,
            size=size,
            weight=ft.FontWeight.BOLD,
            text_align=ft.TextAlign.CENTER,
            color=ft.Colors.BLACK,
        )

    def paragraph(text):
        return ft.Text(text, size=13, color=ft.Colors.BLACK)

    contract = ft.Column(
        width=600,
        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
        controls=[
            title("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", 18),
            ft.Text("Độc lập – Tự do – Hạnh phúc", italic=True),
            ft.Text("———o0o———"),
            title("HỢP ĐỒNG THUÊ XE", 20),
            ft.Text("Số: ……/2026/HĐTX-PiCar"),

            paragraph("- Căn cứ Bộ Luật Dân sự 2015;"),
            paragraph("- Căn cứ Luật Thương mại 2005;"),
            paragraph("- Căn cứ vào nhu cầu và khả năng của các bên."),
            paragraph(""),
            paragraph(
                "Hôm nay, ngày .... tháng .... năm 20..., tại văn phòng PiCar, chúng tôi gồm:"
            ),

            paragraph(""),
            paragraph("BÊN A: BÊN CHO THUÊ (CÔNG TY PiCar)"),
            paragraph("- Địa chỉ: Tầng 1, Tòa nhà Tech, TP.HCM"),
            paragraph("- Đại diện: Ông Nguyễn Văn A"),
            paragraph("- Chức vụ: Giám đốc điều hành"),

            paragraph(""),
            paragraph("BÊN B: BÊN THUÊ (KHÁCH HÀNG)"),
            paragraph("- Ông/Bà: ................................................"),
            paragraph("- Số CCCD/CMND: ........................................."),
            paragraph("- Địa chỉ thường trú: ....................................."),
            paragraph("- Số điện thoại: ........................................"),

            ft.ElevatedButton(
                "XUẤT FILE WORD (.docx)",
                bgcolor=ft.Colors.BLACK,
                color=ft.Colors.WHITE,
                on_click=export_docx,
            ),
        ],
    )

    page.add(ft.Row(alignment=ft.MainAxisAlignment.CENTER, controls=[contract]))


ft.app(target=main)
